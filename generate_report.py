import os
from docx import Document
from docx.shared import Pt, Inches
from docx.enum.text import WD_ALIGN_PARAGRAPH

def generate_report():
    doc = Document()
    
    # helper for centered title text
    def add_centered_line(text, size=12, bold=False, space_after=Pt(0)):
        p = doc.add_paragraph()
        p.alignment = WD_ALIGN_PARAGRAPH.CENTER
        p.paragraph_format.space_after = space_after
        run = p.add_run(text)
        run.font.size = Pt(size)
        run.bold = bold
        return p

    # --- TITLE PAGE ---
    add_centered_line("", space_after=Pt(50))
    add_centered_line("“EXPLORATORY DATA ANALYSIS AND INSIGHT GENERATION", size=14, bold=True)
    add_centered_line("ON THE GLOBAL VIDEO GAME MARKET (1980–2016)”", size=14, bold=True, space_after=Pt(50))
    
    add_centered_line("A PROJECT REPORT SUBMITTED TO BHARATHIAR UNIVERSITY", size=12)
    add_centered_line("IN PARTIAL FULFILMENT OF THE REQUIREMENTS FOR THE", size=12)
    add_centered_line("AWARD OF THE DEGREE OF", size=12, space_after=Pt(30))
    
    add_centered_line("BACHELOR OF DATA SCIENCE", size=14, bold=True, space_after=Pt(50))
    
    add_centered_line("Submitted By", size=12)
    add_centered_line("JAWER-22", size=12, bold=True)
    add_centered_line("(Reg. No. 2328MXXXX)", size=12, space_after=Pt(50)) # Placeholder Reg. No

    # Guidancy part (placeholder as in sample)
    add_centered_line("Under the Guidance of", size=12)
    add_centered_line("PROJECT COORDINATOR", size=12, bold=True)
    add_centered_line("DEPARTMENT OF DATA SCIENCE", size=12, bold=True, space_after=Pt(50))
    
    add_centered_line("MARCH 2026", size=14, bold=True)
    doc.add_page_break()

    # --- CERTIFICATE ---
    doc.add_heading('CERTIFICATE', level=1).alignment = WD_ALIGN_PARAGRAPH.CENTER
    cert_text = (
        "This is to certify that the project work entitled “EXPLORATORY DATA ANALYSIS AND INSIGHT GENERATION "
        "ON THE GLOBAL VIDEO GAME MARKET (1980–2016)”, is a bona-fide work done by JAWER-22 during the "
        "period of 2025-2026 in the Department of Data Science. The project work is an original work "
        "of the candidate and to the best of my knowledge has not been submitted, in part or in full, "
        "for any Diploma / Degree / Associateship / Fellowship or other similar titles in this or any other University."
    )
    doc.add_paragraph(cert_text)
    doc.add_paragraph("\n\nSIGNATURE OF THE GUIDE")
    doc.add_paragraph("\n\nSIGNATURE OF THE HOD\t\t\tSIGNATURE OF THE PRINCIPAL")
    doc.add_page_break()

    # --- ABSTRACT ---
    doc.add_heading('ABSTRACT', level=1).alignment = WD_ALIGN_PARAGRAPH.CENTER
    abstract_text = (
        "The global video game industry has evolved from a niche hobby into a dominant multi-billion dollar entertainment sector. "
        "Understanding market trends, regional preferences, and platform lifecycles is critical for publishers and stakeholders. "
        "The present study focuses on a comprehensive Exploratory Data Analysis (EDA) of the VGChartz dataset, containing "
        "records for over 16,000 video games spanning 1980 to 2016. A systematic data cleaning and feature engineering pipeline "
        "was implemented, addressing missing values and standardizing platform data. Key findings reveal that the market experienced "
        "a significant peak between 2008 and 2009, with North America consistently holding the largest market share (~49%). "
        "Action and Sports genres were identified as high-volume leaders, while Platform and Shooter games showed superior "
        "average sales targets. The results highlight distinct regional taste differences, with Japan uniquely favoring Role-Playing games. "
        "This report demonstrates how data-driven visualizations and statistical analysis can uncover the hidden factors driving "
        "the success of video game titles globally."
    )
    doc.add_paragraph(abstract_text)
    doc.add_page_break()

    # --- CHAPTER 1: INTRODUCTION ---
    doc.add_heading('1.1 INTRODUCTION', level=2)
    intro_text = (
        "The entertainment industry has been revolutionized by the digital gaming era. From early pixelated consoles to high-fidelity "
        "modern hardware, video games have become a massive data-generating engine. This project seeks to explore the patterns "
        "within the global gaming market, focusing on how different factors like Genre, Platform, and Region drive commercial success."
    )
    doc.add_paragraph(intro_text)
    
    doc.add_heading('1.2 PROJECT OBJECTIVE', level=2)
    obj_item1 = doc.add_paragraph("To perform deep exploratory analysis on 16,000+ game sales records.")
    obj_item1.style = 'List Bullet'
    obj_item2 = doc.add_paragraph("To identify regional differences in market behavior (NA, EU, JP).")
    obj_item2.style = 'List Bullet'
    obj_item3 = doc.add_paragraph("To visualize platform lifecycles and the impact of release decade.")
    obj_item3.style = 'List Bullet'
    obj_item4 = doc.add_paragraph("To provide an interactive dashboard for multi-dimensional data exploration.")
    obj_item4.style = 'List Bullet'

    doc.add_heading('1.3 DATA OVERVIEW', level=2)
    data_overview_text = (
        "The dataset includes 11 core attributes for each game, ranging from categorical identifiers to numerical sales totals. "
        "Key features include Rank, Name, Platform, Year, Genre, Publisher, and Regional Sales (NA, EU, JP, Other) in millions."
    )
    doc.add_paragraph(data_overview_text)
    doc.add_page_break()

    # --- CHAPTER 2: SYSTEM STUDY ---
    doc.add_heading('CHAPTER 2: SYSTEM STUDY', level=1)
    doc.add_heading('2.1 SYSTEM SPECIFICATION', level=2)
    
    doc.add_heading('2.1.1 HARDWARE SPECIFICATION', level=3)
    hw_specs = ["Processor: Intel Core i5 or higher", "RAM: 8 GB or more", "Storage: 256 GB SSD", "System Type: 64-bit Operating System"]
    for spec in hw_specs:
        doc.add_paragraph(spec, style='List Bullet')
        
    doc.add_heading('2.1.2 SOFTWARE SPECIFICATION', level=3)
    sw_specs = ["Operating System: Windows 10/11", "Programming Language: Python 3.x", "Environment: Jupyter Notebook / Visual Studio Code", "Key Libraries: Pandas, NumPy, Matplotlib, Seaborn, Streamlit"]
    for spec in sw_specs:
        doc.add_paragraph(spec, style='List Bullet')
    doc.add_page_break()

    # --- CHAPTER 3: SYSTEM DEVELOPMENT ---
    doc.add_heading('CHAPTER 3: SYSTEM DEVELOPMENT', level=1)
    
    doc.add_heading('3.1 DATA PREPROCESSING', level=2)
    prep_text = (
        "Data integrity was ensured by addressing missing values and standardizing columns. "
        "A critical constraint was applied to cap the analysis at the year 2016 to ensure data completeness and recency."
    )
    doc.add_paragraph(prep_text)
    
    doc.add_heading('3.2 EXPLORATORY DATA ANALYSIS (EDA)', level=2)
    eda_text = (
        "EDA uncovers the underlying structure of the global gaming market through distribution analysis and time-series trends."
    )
    doc.add_paragraph(eda_text)

    # Adding the Synchronized Area Chart image
    img_path = r'C:\Users\Nithishvar\.gemini\antigravity\brain\e387e6d1-f5c9-451e-9aff-382c10326f9b\synchronized_area_chart.png'
    if os.path.exists(img_path):
        doc.add_paragraph("Fig 3.1: Regional Sales Trends Over Time (Stacked Area Chart)")
        doc.add_picture(img_path, width=Inches(6))
        doc.paragraphs[-1].alignment = WD_ALIGN_PARAGRAPH.CENTER
    
    doc.add_heading('3.3 KEY INSIGHTS', level=2)
    insights = [
        "2008 Peak: The gaming market reached its highest revenue peak during 2008-2009.",
        "Regional Dominance: North America accounts for nearly half of the global market share.",
        "Japan's Unique Taste: Unlike NA and EU (Shooters/Sports), Japan prefers Role-Playing games."
    ]
    for insight in insights:
        doc.add_paragraph(insight, style='List Bullet')
    doc.add_page_break()

    # --- CHAPTER 4: STREAMLIT WEB INTERFACE ---
    doc.add_heading('CHAPTER 4: STREAMLIT WEB INTERFACE', level=1)
    ui_text = (
        "To provide stakeholders with a self-service tool, an interactive web application was developed using Streamlit. "
        "The interface allows users to filter data by year ranges (1980-2016), genres, and platforms in real-time."
    )
    doc.add_paragraph(ui_text)
    
    # Adding Dashboard screenshot
    dash_path = r'C:\Users\Nithishvar\.gemini\antigravity\brain\e387e6d1-f5c9-451e-9aff-382c10326f9b\main_dashboard_1773369382906.png'
    if os.path.exists(dash_path):
        doc.add_paragraph("Fig 4.1: Interactive Market Analysis Dashboard")
        doc.add_picture(dash_path, width=Inches(6))
        doc.paragraphs[-1].alignment = WD_ALIGN_PARAGRAPH.CENTER
    doc.add_page_break()

    # --- CHAPTER 5: CONCLUSION ---
    doc.add_heading('CHAPTER 5: CONCLUSION', level=1)
    conclusion_text = (
        "The project successfully identifies the primary growth factors of the video game market. "
        "By leveraging automated data cleaning and modern visualization tools, we provide a robust "
        "framework for understanding global gaming commerce."
    )
    doc.add_paragraph(conclusion_text)

    doc.save('PROJECT_REPORT.docx')
    print("Full report generated successfully as PROJECT_REPORT.docx")

if __name__ == "__main__":
    generate_report()
